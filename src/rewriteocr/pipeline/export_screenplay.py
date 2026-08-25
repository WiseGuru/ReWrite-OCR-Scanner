"""Screenplay renderers: Fountain, Final Draft (.fdx), and styled DOCX.

All three are additional renderers over the same canonical Markdown, reached
through the classification pass in core/screenplay.py. The element stream is
built per export and never stored, so there is still exactly one stored
representation and no generation path that bypasses the Markdown.

The existing prose DOCX renderer is untouched; a screenplay DOCX is built
from the element stream directly, because named paragraph styles with fixed
indents have no Markdown token to walk.
"""

from __future__ import annotations

import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from rewriteocr.core.models import (
    ElementType,
    ExportOptions,
    Screenplay,
    StitchLog,
)
from rewriteocr.core.screenplay import SCENE_PREFIX_RE, TITLE_KEYS, parse_screenplay
from rewriteocr.core.screenplay_geom import line_indents
from rewriteocr.core.sidecar import SidecarDB
from rewriteocr.core.stitching import stitch_pages
from rewriteocr.pipeline.export_md import collect_page_texts

FOUNTAIN_PAGE_BREAK = "==="

# Fountain's canonical title-page order. Any other key we detected is
# appended after these, in the order it was found.
TITLE_KEY_ORDER = TITLE_KEYS


def load_screenplay(
    sidecar: SidecarDB, pdf_path: Path | None, options: ExportOptions
) -> tuple[Screenplay, StitchLog]:
    """Canonical Markdown to a classified element stream, plus the export log."""
    pages = collect_page_texts(sidecar)
    if options.stitch:
        pages, log = stitch_pages(list(pages), protect=SCENE_PREFIX_RE)
    else:
        log = StitchLog()
    indents = line_indents(sidecar, pdf_path, pages)
    script = parse_screenplay(pages, indents=indents)
    log.screenplay = script.report
    return script, log


# -- Fountain ---------------------------------------------------------------


def _defuse(text: str) -> str:
    """Neutralize Fountain's span markup in transcribed text. A stray '[[' in
    an OCR'd script would otherwise open a note and swallow everything to the
    next ']]'. Nothing else needs escaping: the forcing characters make every
    element type explicit, so no other convention can fire."""
    return (
        text.replace("[[", "[ [")
        .replace("]]", "] ]")
        .replace("/*", "/ *")
        .replace("*/", "* /")
    )


def _title_block(script: Screenplay) -> list[str]:
    fields = script.title.fields
    if not fields:
        return []
    ordered = [k for k in TITLE_KEY_ORDER if k in fields]
    ordered += [k for k in fields if k not in ordered]
    return [f"{key}: {_defuse(fields[key])}" for key in ordered]


def render_fountain(script: Screenplay, page_break: str = "none") -> str:
    """Emit Fountain using the forcing characters throughout.

    Fountain normally infers element type from convention, and that inference
    misfires: a line of uppercase action reads as a character cue. The
    classifier already determined the type, so it is stated outright rather
    than left for the downstream parser to re-guess.

    The cost, which is the point: centered text and lyrics are never
    auto-detected, and a forced action can never be re-read as a scene
    heading.
    """
    out: list[str] = []
    title = _title_block(script)
    if title:
        out.extend(title)
        out.append("")

    prev_page: int | None = None
    prev_type: ElementType | None = None
    for el in script.elements:
        # A speech is one block: a blank line between a cue and its dialogue
        # would end the speech and orphan the dialogue as action.
        if out and not _in_same_speech(prev_type, el.type):
            out.append("")
        if (
            page_break != "none"
            and prev_page is not None
            and el.page_index != prev_page
        ):
            # A page break is its own block and needs blank lines either side.
            if out and out[-1] != "":
                out.append("")
            out.append(FOUNTAIN_PAGE_BREAK)
            out.append("")

        text = _defuse(el.text)
        if el.type == "scene_heading":
            out.append(f".{text}")
        elif el.type == "action":
            out.append(f"!{text}")
        elif el.type == "character":
            out.append(f"@{text}{' ^' if el.dual else ''}")
        elif el.type == "transition":
            out.append(f"> {text}")
        else:
            # Parenthetical and dialogue have no forcing character in the
            # syntax, and need none: the cue above them is forced with @.
            out.append(text)
        prev_page = el.page_index
        prev_type = el.type

    return "\n".join(out).strip() + "\n"


def _in_same_speech(prev: ElementType | None, current: ElementType) -> bool:
    """A cue, its parentheticals and its dialogue are one block. A
    parenthetical mid-speech follows dialogue, so that pairing counts too."""
    if prev is None:
        return True
    return prev in ("character", "parenthetical", "dialogue") and current in (
        "parenthetical",
        "dialogue",
    )


def export_fountain(
    sidecar: SidecarDB, pdf_path: Path | None, out_path: Path, options: ExportOptions
) -> StitchLog:
    script, log = load_screenplay(sidecar, pdf_path, options)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(
        render_fountain(script, options.page_break), encoding="utf-8", newline="\n"
    )
    return log


# -- Final Draft ------------------------------------------------------------

# Verified against real Final Draft exports, not recollection: the type
# strings, the root attributes, and that <Text> nests directly inside
# <Paragraph> with no wrapper element.
FDX_TYPE: dict[ElementType, str] = {
    "scene_heading": "Scene Heading",
    "action": "Action",
    "character": "Character",
    "parenthetical": "Parenthetical",
    "dialogue": "Dialogue",
    "transition": "Transition",
    "general": "General",
}

FDX_DECLARATION = b'<?xml version="1.0" encoding="UTF-8" standalone="no" ?>\n'


def _take_speech(elements: list, start: int) -> tuple[list, int]:
    """One speech: a cue plus the parentheticals and dialogue under it."""
    speech = [elements[start]]
    i = start + 1
    while i < len(elements) and elements[i].type in ("parenthetical", "dialogue"):
        speech.append(elements[i])
        i += 1
    return speech, i


def dual_groups(elements: list) -> list[tuple[bool, list]]:
    """Group the flat element stream into (is_dual, elements) runs.

    Fountain marks dual dialogue with a caret on the **second** cue, so a
    dual pair is recognized by looking forward from a completed speech to a
    cue carrying `dual`. A dual cue with no speech before it, or either half
    missing its dialogue, degrades to two ordinary speeches rather than
    producing a wrapper Final Draft would reject.
    """
    out: list[tuple[bool, list]] = []
    i = 0
    while i < len(elements):
        if elements[i].type != "character":
            out.append((False, [elements[i]]))
            i += 1
            continue
        first, j = _take_speech(elements, i)
        paired = (
            j < len(elements)
            and elements[j].type == "character"
            and elements[j].dual
        )
        if paired:
            second, k = _take_speech(elements, j)
            if _has_dialogue(first) and _has_dialogue(second):
                out.append((True, first + second))
                i = k
                continue
        out.append((False, first))
        i = j
    return out


def _has_dialogue(speech: list) -> bool:
    return any(el.type == "dialogue" for el in speech)


def _add_paragraph(parent, el) -> None:
    para = ET.SubElement(parent, "Paragraph", {"Type": FDX_TYPE[el.type]})
    ET.SubElement(para, "Text").text = el.text


def render_fdx(script: Screenplay) -> bytes:
    """Emit Final Draft XML.

    Dual dialogue is an untyped wrapper `<Paragraph>` holding a
    `<DualDialogue>` whose children are a **flat** sequence of both speakers,
    left column first:

        <Paragraph>                       <- no Type, no Text of its own
          <DualDialogue>
            <Paragraph Type="Character">
            <Paragraph Type="Dialogue">
            <Paragraph Type="Character">
            <Paragraph Type="Dialogue">
          </DualDialogue>
        </Paragraph>

    Two deliberate omissions remain:

    - **<TitlePage>**, which is not Fountain-style key/value pairs but a
      HeaderAndFooter plus centered paragraphs carrying layout attributes.
      Title fields become leading General paragraphs instead; Final Draft
      opens a script with no title page without complaint.
    - **Forced page breaks.** options.page_break is ignored here: Final Draft
      repaginates on open, so OCR page boundaries would only produce short
      pages.
    """
    root = ET.Element(
        "FinalDraft", {"DocumentType": "Script", "Template": "No", "Version": "3"}
    )
    content = ET.SubElement(root, "Content")

    for key, value in script.title.fields.items():
        para = ET.SubElement(content, "Paragraph", {"Type": "General"})
        ET.SubElement(para, "Text").text = f"{key}: {value}"

    for is_dual, group in dual_groups(script.elements):
        if is_dual:
            wrapper = ET.SubElement(content, "Paragraph")
            dual = ET.SubElement(wrapper, "DualDialogue")
            for el in group:
                _add_paragraph(dual, el)
        else:
            for el in group:
                _add_paragraph(content, el)

    ET.indent(root, space="  ")
    body = ET.tostring(root, encoding="utf-8", xml_declaration=False)
    # ElementTree cannot emit standalone="no", so the declaration is written
    # by hand to match what Final Draft itself produces.
    return FDX_DECLARATION + body + b"\n"


def export_fdx(
    sidecar: SidecarDB, pdf_path: Path | None, out_path: Path, options: ExportOptions
) -> StitchLog:
    script, log = load_screenplay(sidecar, pdf_path, options)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_bytes(render_fdx(script))
    return log


# -- styled DOCX ------------------------------------------------------------


@dataclass(frozen=True)
class _StyleSpec:
    name: str
    left_in: float
    right_in: float
    space_before_pt: int
    caps: bool = False
    keep: bool = False


# Indents are relative to the left margin, which is python-docx's origin for
# left_indent. With the 1.5in left margin set below, that puts dialogue at
# 2.5in from the paper edge, the character cue at 3.5in, the parenthetical at
# 3.0in and the transition at 5.5in, which are the conventional positions.
STYLE_SPECS: dict[ElementType, _StyleSpec] = {
    "scene_heading": _StyleSpec("Scene Heading", 0.0, 0.0, 12, caps=True, keep=True),
    "action": _StyleSpec("Action", 0.0, 0.0, 12),
    "character": _StyleSpec("Character", 2.0, 0.0, 12, caps=True, keep=True),
    "parenthetical": _StyleSpec("Parenthetical", 1.5, 2.0, 0, keep=True),
    "dialogue": _StyleSpec("Dialogue", 1.0, 1.5, 0),
    "transition": _StyleSpec("Transition", 4.0, 0.0, 12, caps=True),
    "general": _StyleSpec("General", 0.0, 0.0, 12),
}

TITLE_STYLE_NAME = "Title Page"
SCREENPLAY_FONT = "Courier New"
SCREENPLAY_FONT_PT = 12


def _set_base_font(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = SCREENPLAY_FONT
    normal.font.size = Pt(SCREENPLAY_FONT_PT)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    # python-docx does not set eastAsia, and Word substitutes a font for
    # anything it does not find a match for there.
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), SCREENPLAY_FONT)


def _ensure_styles(document: Document) -> None:
    """Add the named screenplay paragraph styles. Idempotent: a style that is
    already present (Word ships no 'Scene Heading', but a re-render or a
    future template might) is left alone."""
    from docx.enum.style import WD_STYLE_TYPE

    existing = {s.name for s in document.styles}
    normal = document.styles["Normal"]
    specs = list(STYLE_SPECS.values()) + [_StyleSpec(TITLE_STYLE_NAME, 0.0, 0.0, 0)]
    for spec in specs:
        if spec.name in existing:
            continue
        style = document.styles.add_style(spec.name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.name = SCREENPLAY_FONT
        style.font.size = Pt(SCREENPLAY_FONT_PT)
        if spec.caps:
            style.font.all_caps = True
        pf = style.paragraph_format
        pf.left_indent = Inches(spec.left_in)
        pf.right_indent = Inches(spec.right_in)
        pf.space_before = Pt(spec.space_before_pt)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        pf.keep_with_next = spec.keep


def _set_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)


def build_screenplay_document(script: Screenplay) -> Document:
    """Render the element stream as named paragraph styles with real indents,
    which is what downstream parsers read and how stage plays circulate.

    options.page_break is deliberately not honored: screenplay pagination is
    a property of the format, and forcing OCR page boundaries would produce
    short pages. keep_with_next on headings, cues and parentheticals lets
    Word paginate correctly on its own.
    """
    document = Document()
    _set_page(document)
    _set_base_font(document)
    _ensure_styles(document)

    if script.title.fields:
        ordered = [k for k in TITLE_KEY_ORDER if k in script.title.fields]
        ordered += [k for k in script.title.fields if k not in ordered]
        for key in ordered:
            para = document.add_paragraph(script.title.fields[key], style=TITLE_STYLE_NAME)
            para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    for el in script.elements:
        document.add_paragraph(el.text, style=STYLE_SPECS[el.type].name)
    return document


def export_screenplay_docx(
    sidecar: SidecarDB, pdf_path: Path | None, out_path: Path, options: ExportOptions
) -> StitchLog:
    script, log = load_screenplay(sidecar, pdf_path, options)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    build_screenplay_document(script).save(str(out_path))
    return log
