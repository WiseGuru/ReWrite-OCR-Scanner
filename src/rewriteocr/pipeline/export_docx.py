"""DOCX rendering from canonical Markdown via markdown-it-py tokens.

Headings map to Word heading styles, pipe tables to Word tables, HTML tables
(the table-region output format) to Word tables with merged cells, figures
embed inline, and the page-break option becomes real page breaks.
"""

from __future__ import annotations

import logging
import re
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from markdown_it import MarkdownIt

from rewriteocr.core.models import ExportOptions, StitchLog
from rewriteocr.core.sidecar import SidecarDB
from rewriteocr.core.stitching import stitch_pages
from rewriteocr.pipeline.export_md import collect_page_texts

log = logging.getLogger("rewriteocr.docx")

MAX_LIST_LEVEL = 3
IMAGE_WIDTH_IN = 6.0


class _HtmlTableGrid(HTMLParser):
    """Parses one HTML table into a grid with rowspan/colspan resolved to
    (text, master_cell) placements for python-docx merging."""

    def __init__(self) -> None:
        super().__init__()
        self.rows: list[list[dict]] = []
        self._current_row: list[dict] | None = None
        self._current_cell: dict | None = None

    def handle_starttag(self, tag, attrs):
        attrs = dict(attrs)
        if tag == "tr":
            self._current_row = []
        elif tag in ("td", "th") and self._current_row is not None:
            self._current_cell = {
                "text": "",
                "header": tag == "th",
                "rowspan": max(1, int(attrs.get("rowspan", 1) or 1)),
                "colspan": max(1, int(attrs.get("colspan", 1) or 1)),
            }
        elif tag == "br" and self._current_cell is not None:
            self._current_cell["text"] += "\n"

    def handle_endtag(self, tag):
        if tag in ("td", "th") and self._current_cell is not None:
            if self._current_row is not None:
                self._current_row.append(self._current_cell)
            self._current_cell = None
        elif tag == "tr" and self._current_row is not None:
            self.rows.append(self._current_row)
            self._current_row = None

    def handle_data(self, data):
        if self._current_cell is not None:
            self._current_cell["text"] += data

    def grid(self) -> list[list[dict | None]]:
        """Expand spans into a rectangular grid. Each occupied slot holds
        {'text', 'header', 'master': (r, c)}; continuation slots point at
        their master."""
        if not self.rows:
            return []
        n_cols = 0
        for row in self.rows:
            n_cols = max(n_cols, sum(c["colspan"] for c in row))
        grid: list[list[dict | None]] = []
        pending: dict[tuple[int, int], dict] = {}
        for r, row in enumerate(self.rows):
            grid.append([None] * n_cols)
            c = 0
            cells = iter(row)
            while c < n_cols:
                if (r, c) in pending:
                    grid[r][c] = pending.pop((r, c))
                    c += 1
                    continue
                cell = next(cells, None)
                if cell is None:
                    break
                master = {"text": cell["text"].strip(), "header": cell["header"], "master": (r, c)}
                for dr in range(cell["rowspan"]):
                    for dc in range(cell["colspan"]):
                        slot = {"text": "", "header": cell["header"], "master": (r, c)}
                        if dr == 0 and dc == 0:
                            slot = master
                        if dr == 0:
                            if c + dc < n_cols:
                                grid[r][c + dc] = slot
                        else:
                            pending[(r + dr, c + dc)] = slot
                c += cell["colspan"]
        return grid


class DocxRenderer:
    """Walks markdown-it tokens and emits python-docx content."""

    def __init__(self, document: Document, base_dir: Path) -> None:
        self.doc = document
        self.base_dir = base_dir
        self.md = MarkdownIt("commonmark").enable(["table", "strikethrough"])

    def render_markdown(self, markdown: str) -> None:
        tokens = self.md.parse(markdown)
        self._walk(tokens)

    # -- block walker --------------------------------------------------------

    def _walk(self, tokens) -> None:
        i = 0
        list_stack: list[str] = []
        while i < len(tokens):
            tok = tokens[i]
            t = tok.type
            if t == "heading_open":
                level = int(tok.tag[1])
                inline = tokens[i + 1]
                para = self.doc.add_heading("", level=min(level, 9))
                self._inline(para, inline)
                i += 3
            elif t == "paragraph_open":
                inline = tokens[i + 1]
                style = None
                if list_stack:
                    depth = min(len(list_stack), MAX_LIST_LEVEL)
                    base = "List Bullet" if list_stack[-1] == "bullet" else "List Number"
                    style = base if depth == 1 else f"{base} {depth}"
                para = self.doc.add_paragraph(style=style)
                self._inline(para, inline)
                i += 3
            elif t == "bullet_list_open":
                list_stack.append("bullet")
                i += 1
            elif t == "ordered_list_open":
                list_stack.append("ordered")
                i += 1
            elif t in ("bullet_list_close", "ordered_list_close"):
                if list_stack:
                    list_stack.pop()
                i += 1
            elif t == "blockquote_open":
                close = self._find_close(tokens, i, "blockquote_close")
                for j in range(i + 1, close):
                    if tokens[j].type == "inline":
                        para = self.doc.add_paragraph(style="Quote")
                        self._inline(para, tokens[j])
                i = close + 1
            elif t in ("fence", "code_block"):
                para = self.doc.add_paragraph()
                run = para.add_run(tok.content.rstrip("\n"))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                i += 1
            elif t == "table_open":
                close = self._find_close(tokens, i, "table_close")
                self._pipe_table(tokens[i:close + 1])
                i = close + 1
            elif t == "html_block":
                if "<table" in tok.content.lower():
                    self._html_table(tok.content)
                else:
                    text = re.sub(r"<[^>]+>", "", tok.content).strip()
                    if text:
                        self.doc.add_paragraph(text)
                i += 1
            elif t == "hr":
                para = self.doc.add_paragraph()
                para.add_run().add_break(WD_BREAK.PAGE)
                i += 1
            else:
                i += 1

    @staticmethod
    def _find_close(tokens, start: int, close_type: str) -> int:
        depth = 0
        open_type = close_type.replace("_close", "_open")
        for j in range(start, len(tokens)):
            if tokens[j].type == open_type:
                depth += 1
            elif tokens[j].type == close_type:
                depth -= 1
                if depth == 0:
                    return j
        return len(tokens) - 1

    # -- inline rendering ----------------------------------------------------

    def _inline(self, para, inline_token) -> None:
        bold = italic = code = strike = False
        for child in inline_token.children or []:
            t = child.type
            if t == "text":
                run = para.add_run(child.content)
                run.bold, run.italic, run.font.strike = bold or None, italic or None, strike or None
                if code:
                    run.font.name = "Consolas"
            elif t == "code_inline":
                run = para.add_run(child.content)
                run.font.name = "Consolas"
            elif t == "strong_open":
                bold = True
            elif t == "strong_close":
                bold = False
            elif t == "em_open":
                italic = True
            elif t == "em_close":
                italic = False
            elif t == "s_open":
                strike = True
            elif t == "s_close":
                strike = False
            elif t in ("softbreak", "hardbreak"):
                para.add_run().add_break()
            elif t == "image":
                self._image(para, child)
            elif t == "link_open":
                pass  # link text renders as plain text; URL appended on close
            elif t == "link_close":
                pass

    def _image(self, para, token) -> None:
        src = token.attrGet("src") or ""
        path = (self.base_dir / src).resolve()
        if path.is_file():
            try:
                run = para.add_run()
                run.add_picture(str(path), width=Inches(IMAGE_WIDTH_IN))
                return
            except (OSError, ValueError) as exc:
                log.warning("Could not embed image %s: %s", path, exc)
        alt = token.content or "figure"
        para.add_run(f"[{alt}: {src}]")

    # -- tables --------------------------------------------------------------

    def _pipe_table(self, tokens) -> None:
        rows: list[tuple[bool, list]] = []
        in_header = False
        current: list | None = None
        for tok in tokens:
            if tok.type == "thead_open":
                in_header = True
            elif tok.type == "thead_close":
                in_header = False
            elif tok.type == "tr_open":
                current = []
            elif tok.type == "tr_close" and current is not None:
                rows.append((in_header, current))
                current = None
            elif tok.type == "inline" and current is not None:
                current.append(tok)
        if not rows:
            return
        n_cols = max(len(cells) for _, cells in rows)
        table = self.doc.add_table(rows=len(rows), cols=n_cols)
        table.style = "Table Grid"
        for r, (is_header, cells) in enumerate(rows):
            for c, inline in enumerate(cells):
                cell = table.cell(r, c)
                cell.text = ""
                para = cell.paragraphs[0]
                self._inline(para, inline)
                if is_header:
                    for run in para.runs:
                        run.bold = True

    def _html_table(self, html: str) -> None:
        parser = _HtmlTableGrid()
        try:
            parser.feed(html)
        except Exception as exc:
            log.warning("HTML table parse failed: %s", exc)
            self.doc.add_paragraph(re.sub(r"<[^>]+>", " ", html).strip())
            return
        grid = parser.grid()
        if not grid:
            return
        n_rows, n_cols = len(grid), len(grid[0])
        table = self.doc.add_table(rows=n_rows, cols=n_cols)
        table.style = "Table Grid"
        done_masters: set[tuple[int, int]] = set()
        for r in range(n_rows):
            for c in range(n_cols):
                slot = grid[r][c]
                if slot is None:
                    continue
                mr, mc = slot["master"]
                if (mr, mc) in done_masters or (r, c) != (mr, mc):
                    continue
                done_masters.add((mr, mc))
                # Find the span extent of this master in the grid.
                max_r, max_c = mr, mc
                for rr in range(mr, n_rows):
                    if grid[rr][mc] and grid[rr][mc]["master"] == (mr, mc):
                        max_r = rr
                for cc in range(mc, n_cols):
                    if grid[mr][cc] and grid[mr][cc]["master"] == (mr, mc):
                        max_c = cc
                cell = table.cell(mr, mc)
                if (max_r, max_c) != (mr, mc):
                    cell = cell.merge(table.cell(max_r, max_c))
                cell.text = slot["text"]
                if slot["header"]:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True


def export_docx(
    sidecar: SidecarDB, out_path: Path, options: ExportOptions
) -> StitchLog:
    pages = collect_page_texts(sidecar)
    if options.stitch:
        pages, stitch_log = stitch_pages(pages)
    else:
        stitch_log = StitchLog()

    document = Document()
    renderer = DocxRenderer(document, base_dir=sidecar.path.parent)
    first = True
    for text in pages:
        if not text.strip():
            continue
        if not first and options.page_break != "none":
            para = document.add_paragraph()
            para.add_run().add_break(WD_BREAK.PAGE)
        renderer.render_markdown(text)
        first = False
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))
    return stitch_log
