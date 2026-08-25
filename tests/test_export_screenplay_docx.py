"""Styled screenplay DOCX: named paragraph styles with real indents.

Downstream parsers read the style names, which is what makes this different
from the prose DOCX exporter and why stage plays circulate this way.
"""

from __future__ import annotations

import docx
import pytest
from docx.shared import Inches, Pt

from rewriteocr.core.models import ExportOptions, PageRecord
from rewriteocr.core.screenplay import parse_screenplay
from rewriteocr.core.sidecar import SidecarDB
from rewriteocr.pipeline.export_screenplay import (
    build_screenplay_document,
    export_screenplay_docx,
)

SCRIPT_PAGES = [
    "FADE IN:\n\nINT. KITCHEN - DAY\n\nMARLOWE stands at the window.\n\n"
    "MARLOWE\n(quietly)\nYou should not have come here.\n\nCUT TO:",
]


@pytest.fixture
def sidecar(tmp_path):
    db = SidecarDB(tmp_path / "d.ocrproj")
    db.initialize("h", "d.pdf", 1)
    db.insert_pages(
        [PageRecord(page_index=0, classification="scanned", width_pt=612, height_pt=792)]
    )
    db.write_page_result(0, SCRIPT_PAGES[0], "vlm", None, None, [])
    yield db
    db.close()


@pytest.fixture
def rendered(tmp_path):
    out = tmp_path / "script.docx"
    build_screenplay_document(parse_screenplay(SCRIPT_PAGES)).save(str(out))
    return docx.Document(str(out))


def test_named_styles_exist(rendered):
    names = {s.name for s in rendered.styles}
    for name in (
        "Scene Heading",
        "Action",
        "Character",
        "Parenthetical",
        "Dialogue",
        "Transition",
        "General",
    ):
        assert name in names, name


def test_style_indents(rendered):
    # Relative to the 1.5in left margin, so from the paper edge these are
    # dialogue 2.5in, parenthetical 3.0in, character 3.5in, transition 5.5in.
    styles = rendered.styles
    assert styles["Action"].paragraph_format.left_indent == Inches(0.0)
    assert styles["Dialogue"].paragraph_format.left_indent == Inches(1.0)
    assert styles["Dialogue"].paragraph_format.right_indent == Inches(1.5)
    assert styles["Parenthetical"].paragraph_format.left_indent == Inches(1.5)
    assert styles["Character"].paragraph_format.left_indent == Inches(2.0)
    assert styles["Transition"].paragraph_format.left_indent == Inches(4.0)


def test_paragraph_styles_are_assigned(rendered):
    pairs = [(p.style.name, p.text) for p in rendered.paragraphs if p.text]
    assert ("Action", "FADE IN:") in pairs
    assert ("Scene Heading", "INT. KITCHEN - DAY") in pairs
    assert ("Character", "MARLOWE") in pairs
    assert ("Parenthetical", "(quietly)") in pairs
    assert ("Dialogue", "You should not have come here.") in pairs
    assert ("Transition", "CUT TO:") in pairs


def test_page_setup(rendered):
    section = rendered.sections[0]
    assert section.page_width == Inches(8.5)
    assert section.page_height == Inches(11)
    assert section.left_margin == Inches(1.5)
    assert section.right_margin == Inches(1.0)


def test_base_font_is_courier(rendered):
    normal = rendered.styles["Normal"]
    assert normal.font.name == "Courier New"
    assert normal.font.size == Pt(12)


def test_caps_and_keep_with_next_are_set(rendered):
    styles = rendered.styles
    assert styles["Character"].font.all_caps is True
    assert styles["Scene Heading"].font.all_caps is True
    assert styles["Dialogue"].font.all_caps in (None, False)
    assert styles["Character"].paragraph_format.keep_with_next is True
    assert styles["Scene Heading"].paragraph_format.keep_with_next is True


def test_title_page_is_centered_and_followed_by_a_break(tmp_path):
    pages = ["Title: The Long Goodbye\nAuthor: Jane Doe", "INT. KITCHEN - DAY"]
    out = tmp_path / "titled.docx"
    build_screenplay_document(parse_screenplay(pages)).save(str(out))
    doc = docx.Document(str(out))
    first = doc.paragraphs[0]
    assert first.style.name == "Title Page"
    assert first.text == "The Long Goodbye"
    assert 'w:type="page"' in doc.element.xml


def test_export_writes_the_file(sidecar, tmp_path):
    out = tmp_path / "script.docx"
    log = export_screenplay_docx(
        sidecar, None, out, ExportOptions(fmt="screenplay_docx", stitch=False)
    )
    doc = docx.Document(str(out))
    assert ("Character", "MARLOWE") in [
        (p.style.name, p.text) for p in doc.paragraphs if p.text
    ]
    assert log.screenplay is not None


def test_prose_docx_exporter_is_untouched(tmp_path):
    """The Markdown-driven DOCX renderer must keep its own behavior; this is
    a separate renderer over the same canonical Markdown, not a replacement."""
    from rewriteocr.pipeline.export_docx import export_docx

    db = SidecarDB(tmp_path / "prose.ocrproj")
    db.initialize("h", "p.pdf", 1)
    db.insert_pages(
        [PageRecord(page_index=0, classification="scanned", width_pt=612, height_pt=792)]
    )
    db.write_page_result(0, "# Title\n\nBody text.", "vlm", None, None, [])
    out = tmp_path / "prose.docx"
    export_docx(db, out, ExportOptions(fmt="docx", stitch=False))
    db.close()
    styles = [(p.style.name, p.text) for p in docx.Document(str(out)).paragraphs if p.text]
    assert ("Heading 1", "Title") in styles
    assert "Scene Heading" not in {name for name, _ in styles}
