import docx
import pytest

from rewriteocr.core.models import ExportOptions, PageRecord
from rewriteocr.core.sidecar import SidecarDB
from rewriteocr.pipeline.export_docx import _HtmlTableGrid, export_docx


@pytest.fixture
def sidecar(tmp_path):
    db = SidecarDB(tmp_path / "d.ocrproj")
    db.initialize("h", "d.pdf", 2)
    db.insert_pages(
        [PageRecord(page_index=i, classification="scanned", width_pt=612, height_pt=792)
         for i in range(2)]
    )
    yield db
    db.close()


def _texts(doc):
    return [p.text for p in doc.paragraphs]


def test_headings_lists_and_emphasis(sidecar, tmp_path):
    sidecar.write_page_result(
        0,
        "# Title\n\n## Section\n\nSome **bold** and *italic* text.\n\n"
        "- item one\n- item two\n\n1. first\n2. second",
        "vlm:m", "m", "r", [],
    )
    out = tmp_path / "out.docx"
    export_docx(sidecar, out, ExportOptions(fmt="docx", stitch=False))
    doc = docx.Document(str(out))
    styles = [(p.style.name, p.text) for p in doc.paragraphs if p.text]
    assert ("Heading 1", "Title") in styles
    assert ("Heading 2", "Section") in styles
    assert any(s.startswith("List Bullet") and t == "item one" for s, t in styles)
    assert any(s.startswith("List Number") and t == "first" for s, t in styles)
    body = next(p for p in doc.paragraphs if "bold" in p.text)
    bold_runs = [r.text for r in body.runs if r.bold]
    italic_runs = [r.text for r in body.runs if r.italic]
    assert "bold" in "".join(bold_runs)
    assert "italic" in "".join(italic_runs)


def test_pipe_table_renders_as_word_table(sidecar, tmp_path):
    sidecar.write_page_result(
        0, "| A | B |\n| --- | --- |\n| 1 | 2 |", "vlm:m", "m", "r", []
    )
    out = tmp_path / "t.docx"
    export_docx(sidecar, out, ExportOptions(fmt="docx", stitch=False))
    doc = docx.Document(str(out))
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.cell(0, 0).text == "A"
    assert table.cell(1, 1).text == "2"
    assert table.cell(0, 0).paragraphs[0].runs[0].bold


def test_html_table_with_merged_cells(sidecar, tmp_path):
    html = (
        '<table><tr><th colspan="2">Header</th></tr>'
        "<tr><td>a</td><td>b</td></tr>"
        '<tr><td rowspan="2">tall</td><td>c</td></tr>'
        "<tr><td>d</td></tr></table>"
    )
    sidecar.write_page_result(0, html, "vlm:m", "m", "r", [])
    out = tmp_path / "h.docx"
    export_docx(sidecar, out, ExportOptions(fmt="docx", stitch=False))
    doc = docx.Document(str(out))
    assert len(doc.tables) == 1
    table = doc.tables[0]
    # Merged header spans both columns: same underlying cell object.
    assert table.cell(0, 0)._tc is table.cell(0, 1)._tc
    assert table.cell(0, 0).text == "Header"
    assert table.cell(2, 0)._tc is table.cell(3, 0)._tc
    assert table.cell(2, 0).text == "tall"
    assert table.cell(3, 1).text == "d"


def test_page_break_between_pages(sidecar, tmp_path):
    sidecar.write_page_result(0, "Page one.", "vlm:m", "m", "r", [])
    sidecar.write_page_result(1, "Page two.", "vlm:m", "m", "r", [])
    out = tmp_path / "pb.docx"
    export_docx(sidecar, out, ExportOptions(fmt="docx", page_break="comment", stitch=False))
    doc = docx.Document(str(out))
    xml = doc.element.xml
    assert 'w:type="page"' in xml


def test_missing_figure_degrades_to_placeholder(sidecar, tmp_path):
    sidecar.write_page_result(0, "![Figure](nope_figures/gone.png)", "vlm:m", "m", "r", [])
    out = tmp_path / "fig.docx"
    export_docx(sidecar, out, ExportOptions(fmt="docx", stitch=False))
    doc = docx.Document(str(out))
    assert any("gone.png" in p.text for p in doc.paragraphs)


def test_html_grid_expansion():
    parser = _HtmlTableGrid()
    parser.feed('<table><tr><td rowspan="2">x</td><td>a</td></tr><tr><td>b</td></tr></table>')
    grid = parser.grid()
    assert len(grid) == 2 and len(grid[0]) == 2
    assert grid[0][0]["master"] == (0, 0)
    assert grid[1][0]["master"] == (0, 0)
    assert grid[1][1]["text"] == "b"
